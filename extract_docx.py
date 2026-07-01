#!/usr/bin/env python3
"""
Extract text and tables from FIFA World Cup DOCX file
Output: text_extracted.txt with structured content
"""

from docx import Document
from docx.table import Table
import os
import json
def extract_tables_from_docx(docx_path):
    """Extract all tables from DOCX"""
    doc = Document(docx_path)
    tables_data = []
    
    for table_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(cells)
        
        tables_data.append({
            "table_id": table_idx,
            "rows": rows
        })
    
    return tables_data

def extract_text_from_docx(docx_path):
    """Extract all paragraphs from DOCX with structure"""
    doc = Document(docx_path)
    content = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text:  # Skip empty paragraphs
            # Detect heading level
            heading_level = None
            if paragraph.style and paragraph.style.name and paragraph.style.name.startswith('Heading'):
                heading_level = paragraph.style.name.replace('Heading ', '')
            
            content.append({
                "para_id": para_idx,
                "text": text,
                "heading_level": heading_level,
                "is_heading": heading_level is not None
            })
    
    return content

def format_table_as_text(table_data):
    """Convert table dict to readable text"""
    text = ""
    rows = table_data["rows"]
    
    # Header row
    if rows:
        header = " | ".join(rows[0])
        text += header + "\n"
        text += "-" * len(header) + "\n"
        
        # Data rows
        for row in rows[1:]:
            text += " | ".join(row) + "\n"
    
    return text

def main():
    docx_path=r"C:\Users\E_Magic\Downloads\Fifa world cup\FIFA_World_Cup_Professional_Guide.docx"

    
    if not os.path.exists(docx_path):
        print(f"❌ Error: File not found at {docx_path}")
        return
    
    print("📄 Extracting text from DOCX...")
    paragraphs = extract_text_from_docx(docx_path)
    
    print("📊 Extracting tables from DOCX...")
    tables = extract_tables_from_docx(docx_path)
    
    # Save extracted content as structured text
    output_text = "./data/text_extracted.txt"
    os.makedirs("./data", exist_ok=True)
    
    with open(output_text, 'w', encoding='utf-8') as f:
        f.write("=" * 80 + "\n")
        f.write("FIFA WORLD CUP - EXTRACTED TEXT & TABLES\n")
        f.write("=" * 80 + "\n\n")
        
        # Write all paragraphs
        for para in paragraphs:
            if para["is_heading"]:
                f.write(f"\n### HEADING {para['heading_level']}: {para['text']}\n\n")
            else:
                f.write(f"{para['text']}\n\n")
        
        # Write all tables
        f.write("\n" + "=" * 80 + "\n")
        f.write("EXTRACTED TABLES\n")
        f.write("=" * 80 + "\n\n")
        
        for table in tables:
            f.write(f"--- TABLE {table['table_id']} ---\n")
            f.write(format_table_as_text(table))
            f.write("\n\n")
    
    # Save structured JSON
    output_json ="./data/extracted_structure.json" 
    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump({
            "paragraphs": paragraphs,
            "tables": tables,
            "total_paragraphs": len(paragraphs),
            "total_tables": len(tables)
        }, f, indent=2, ensure_ascii=False)
    
    # Statistics
    total_words = sum(len(p["text"].split()) for p in paragraphs)
    
    print(f"""
✅ Extraction Complete!

📊 Statistics:
   - Total Paragraphs: {len(paragraphs)}
   - Total Tables: {len(tables)}
   - Total Words: {total_words}
   
📁 Output Files:
   - {output_text}
   - {output_json}
    """)
    
    # Preview
    print("\n📋 Preview (First 5 paragraphs):")
    for i, para in enumerate(paragraphs[:5]):
        if para["is_heading"]:
            print(f"   [{i}] HEADING: {para['text'][:60]}...")
        else:
            print(f"   [{i}] {para['text'][:60]}...")

if __name__ == "__main__":
    main()
